import docx
from googletrans import Translator  # try pip install googletrans==4.0.0-rc1, if there is a Module Not Found Error
import os
import platform


def slash():
    operating_system = platform.system()
    if operating_system == 'Windows':
        return "\\"
    else:
        return '/'


def get_path(file):
    print("your current working directory is", os.getcwd())

    while True:
        change_cwd = input("Do you want to change directory enter Y or N ::").upper()
        if change_cwd == "y" or "N":
            break
        else:
            print(f"you have entered {change_cwd} which is a incorrect value")

    if change_cwd == "Y":
        new_cwd = input("enter the new directory ::")
        os.chdir(new_cwd)
    file[0] = os.getcwd()
    file[1] = slash()
    file[2] = input("enter the file name in file_name.docx format ::")
    return file


translator = Translator()

n = int(input("enter the number of documents you want to translate "))
docx_list = list(range(n))

for r in range(len(docx_list)):
    docx_list[r] = list((range(3)))

for file_path in range(n):
    docx_list[file_path] = get_path(docx_list[file_path])

for file_path in docx_list:

    ms_word = docx.Document(file_path[0] + file_path[1] + file_path[2])

    paragraphs = ms_word.paragraphs
    no_paragraphs = len(paragraphs)

    translated_word = docx.Document()

    for paragraph_index in range(no_paragraphs):

        paragraph_lines = ms_word.paragraphs[paragraph_index].text
        str_len_not_0 = True
        print(paragraph_lines)

        if len(paragraph_lines) == 0:
            str_len_not_0 = False

        if paragraph_lines is not None and str_len_not_0:  # to avoid type error during conversion

            translation_result = translator.translate(paragraph_lines)  # note : since source language can be
            # auto-detected we don't need to mention src=th, by default dest= english, hence we don't need to declare it
            print(translation_result.text)
            translated_word.add_paragraph(translation_result.text)

    translated_word.save(file_path[0] + file_path[1] + 'translated' + file_path[2])  # it must be saved to unique name
    # within its file path as MS Word would not give permission to replace an existing Word docx
