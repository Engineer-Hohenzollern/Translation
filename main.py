import docx
from googletrans import Translator  # try pip install googletrans==4.0.0-rc1, if there is a Module Not Found Error

translator = Translator()

n = int(input("enter the number of documents you want to translate "))
l = list(range(n))
for r in range(n):
    l[r] = input("enter the file path")

for h in l:
    d = docx.Document(h)  # to open the document containing the promotional content
    i = d.paragraphs  # i return's a list of object where each of these object stores the paragraph elements as strings
    #  we can access the string stored in these objects by d.paragraphs[list_element].text

    i = len(i)  # inorder to get the number of paragraphs in word document, so we can iterate through d.paragraphs

    newd = docx.Document()  # we're creating a new Word document for the translated text
    for r in range(i):
        c = d.paragraphs[r].text  # each paragraphs of the Word document is  returned as a string
        a = True
        print(c)
        if len(c) == 0:
            a = False
        if c is not None and a:  # we don't want empty lines to be translated otherwise this might cause a type error
            result = translator.translate(
                c)  # here  we translate it from thai to eng , note : since source language can be
            # auto-detected we don't need to mention src=th, by default dest= english, hence we don't need to declare it
            print(result.text)
            newd.add_paragraph(result.text)  # we add the translated paragraphs to the new Word document

    newd.save('trans' + h)  # it must be saved to unique name within its file path as MS Word would not give
    # permission to replace an existing Word docx
