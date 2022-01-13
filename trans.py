from googletrans import Translator  # try pip install googletrans==4.0.0-rc1, if there is a Module Not Found Error
import os
import docx
import platform
import argparse
from gtts import gTTS


def tts(text, path):
    tts = gTTS(text)
    tts.save(path + '.mp3')


def add_string(add_str):
    text = add_str + '\n'
    return text


# TODO: Check if os.path.join detects OS "yes it dose i will change it where ever possible"

def slash():
    operating_system = platform.system()
    if operating_system == 'Windows':
        return "\\"
    else:
        return '/'

def pdf_translation(all_pdf):
    pass

def word_translation(all_word_Docx):
    print(all_word_Docx)
    file_len = len(all_word_Docx)

    docx_list = list(range(file_len))
    translated_files_path = os.path.join(path, file + 'translated_files')
    os.mkdir(translated_files_path)

    for r in range(file_len):
        docx_list[r] = list((range(3)))

    for file_path in range(file_len):
        docx_list[file_path] = [path, slash(), all_word_Docx[file_path]]

    for file_path in docx_list:

        ms_word = docx.Document(file_path[0] + file_path[1] + file_path[2])

        paragraphs = ms_word.paragraphs
        no_paragraphs = len(paragraphs)

        translated_word = docx.Document()
        string = str()

        for paragraph_index in range(no_paragraphs):

            paragraph_lines = ms_word.paragraphs[paragraph_index].text
            str_len_not_0 = True
            # print(paragraph_lines)

            if len(paragraph_lines) == 0:
                str_len_not_0 = False

            if paragraph_lines is not None and str_len_not_0:  # to avoid type error during conversion

                translation_result = translator.translate(paragraph_lines)  # note : since source language can be
                # auto-detected we don't need to mention src=th, by default dest= english, hence we don't need to
                # declare it print(translation_result.text)
                string = add_string(translation_result.text)
                translated_word.add_paragraph(translation_result.text)

        os.chdir(translated_files_path)
        translated_word.save(os.getcwd() + file_path[1] + 'trans_' + file_path[2])  # it must be saved to
        #  a unique name within its file path as MS Word would not give permission to replace an existing Word docx
        tts(string, os.getcwd() + file_path[1] + 'trans_' + file_path[2].split('.')[0])
        # os.chdir(path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=f" your current path is {os.getcwd()} ")

    # TODO: Change the file and paths they can be mutually exclusive  "Done"
    parser.add_argument('path', metavar='path', type=str, help='enter the  path')
    parser.add_argument('file', metavar='file', type=str, help='enter the file name')
    args = parser.parse_args()
    path = args.path
    file = args.file
    translator = Translator()

    if file != "*":
        all_documents = os.listdir(os.path.join(path, file))
        all_pdf = [pdf for pdf in all_documents if pdf.endswith('.pdf')]
        all_word_Docx = [word for word in all_documents if word.endswith('.docx')]
        all_png = [png for png in all_documents if png.endswith('.png')]
    else:
        file=str()
        all_documents = os.listdir(path)
        all_pdf = [pdf for pdf in all_documents if pdf.endswith('.pdf')]
        all_word_Docx = [word for word in all_documents if word.endswith('.docx')]
        all_png = [png for png in all_documents if png.endswith('.png')]

    word_translation(all_word_Docx)
    pdf_translation(all_pdf)

