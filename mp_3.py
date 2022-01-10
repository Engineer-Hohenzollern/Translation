import docx
from googletrans import Translator  # try pip install googletrans==4.0.0-rc1, if there is a Module Not Found Error
import os
import platform
import argparse
from gtts import gTTS


def tts(text,path):
    tts = gTTS(text)
    tts.save(path+'.mp3')


def add_string(add_str):
    text = add_str + '\n'
    return text


parser = argparse.ArgumentParser(description=f" your current path is {os.getcwd()} ")
parser.add_argument('path', metavar='path', type=str, help='enter the file path')
parser.add_argument('file', metavar='file', type=str, help='enter the file path')
args = parser.parse_args()
path = args.path
file = args.file


def slash():
    operating_system = platform.system()
    if operating_system == 'Windows':
        return "\\"
    else:
        return '/'


def get_path(file):
    print("your current working directory is", os.getcwd())

    while True:
        change_cwd = input("Do you want to change directory enter Y or N ::").upper()
        if change_cwd == "y" or "N":
            break
        else:
            print(f"you have entered {change_cwd} which is a incorrect value")

    if change_cwd == "Y":
        new_cwd = input("enter the new directory ::")
        os.chdir(new_cwd)
    file[0] = os.getcwd()
    file[1] = slash()
    file[2] = input("enter the file name in file_name.docx format ::")
    return file


translator = Translator()

n = int(input("enter the number of documents you want to translate including the initial document ::"))
docx_list = list(range(n))

for r in range(len(docx_list)):
    docx_list[r] = list((range(3)))

a = 1
for file_path in range(n):
    if a == 1:
        docx_list[file_path] = [path, slash(), file]
    else:
        docx_list[file_path] = get_path(docx_list[file_path])
    a += 1

for file_path in docx_list:

    ms_word = docx.Document(file_path[0] + file_path[1] + file_path[2])

    paragraphs = ms_word.paragraphs
    no_paragraphs = len(paragraphs)

    translated_word = docx.Document()
    string = str()

    for paragraph_index in range(no_paragraphs):

        paragraph_lines = ms_word.paragraphs[paragraph_index].text
        str_len_not_0 = True
        # print(paragraph_lines)

        if len(paragraph_lines) == 0:
            str_len_not_0 = False

        if paragraph_lines is not None and str_len_not_0:  # to avoid type error during conversion

            translation_result = translator.translate(paragraph_lines)  # note : since source language can be
            # auto-detected we don't need to mention src=th, by default dest= english, hence we don't need to declare it
            # print(translation_result.text)
            string = add_string(translation_result.text)
            translated_word.add_paragraph(translation_result.text)

    translated_word.save(file_path[0] + file_path[1] + 'translated' + file_path[2])  # it must be saved to unique name
    # within its file path as MS Word would not give permission to replace an existing Word docx
    tts(string,file_path[0] + file_path[1] + 'translated' + file_path[2])

